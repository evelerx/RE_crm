from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "Northstone-Copyright-Notice.docx"


SECTIONS = [
    (
        "Ownership of the Platform",
        [
            "All rights, title, and interest in and to Northstone CRM, Deal Intelligence OS, and all related software, system architecture, workflows, features, interfaces, dashboards, builder-site tools, communication modules, integrations, APIs, analytics outputs, documents, branding, and associated intellectual property are owned exclusively by Northstone CRM and/or its licensors unless expressly stated otherwise in writing."
        ],
    ),
    (
        "Materials Covered by Copyright",
        [
            "Source code, object code, scripts, and compiled assets.",
            "Backend logic, API structures, request/response schemas, and route systems.",
            "Frontend components, layouts, forms, dashboards, and UI flows.",
            "Database structures, schema design, data models, and protected arrangement of information.",
            "Builder website templates, rendering systems, publishing workflows, and listing layouts.",
            "Generated documents, reports, analytics views, chart layouts, and formatted output systems.",
            "Text, graphics, icons, screenshots, visual identity assets, and platform presentation materials.",
        ],
    ),
    (
        "Protected Platform Functionality",
        [
            "Northstone CRM claims copyright in the specific software implementation, structure, arrangement, design, wording, formatting, and visual or technical expression of its CRM features, including but not limited to:",
            "Lead, contact, deal, and pipeline management workflows.",
            "Activity tracking, follow-up systems, and communication workflows.",
            "Enterprise owner, employee, and admin control systems.",
            "Builder document generation and builder website publishing workflows.",
            "Property listing, property image, and property presentation systems.",
            "AI-assisted content generation, formatting logic, summaries, and builder copy systems.",
            "Analytics, ROI, insights, security, support, and subscription visibility systems.",
        ],
    ),
    (
        "User Data and Customer Content",
        [
            "Users retain ownership of the business information they enter into the CRM, including contacts, notes, deal data, property content, uploaded images, and related records, subject to the rights necessary for Northstone CRM to host, process, transmit, display, secure, and support the service.",
            "Northstone CRM retains ownership of the platform, software, templates, rendering systems, builder-site framework, and all underlying technology.",
        ],
    ),
    (
        "AI-Assisted and System-Generated Output",
        [
            "Any AI-assisted, rule-based, or system-generated output produced through the platform, including drafts, follow-up content, website copy, builder documents, summaries, and formatted CRM output, is generated using proprietary or licensed systems used by Northstone CRM.",
            "Use of that output does not transfer ownership of the underlying platform, software, prompt structures, formatting logic, or generation systems.",
        ],
    ),
    (
        "Builder Website and Template Rights",
        [
            "Any builder website, listing page, or microsite generated through Northstone CRM is created using proprietary software systems and templates owned by Northstone CRM.",
            "Users may use those sites within their subscription rights, but do not acquire ownership of the underlying builder engine, page structures, reusable templates, rendering logic, or publishing framework.",
        ],
    ),
    (
        "Restrictions",
        [
            "Except where expressly permitted in writing, no person or entity may:",
            "Copy, reproduce, mirror, republish, adapt, translate, modify, or create derivative works from the platform.",
            "Reverse engineer, decompile, disassemble, extract, or attempt to recreate the software or templates.",
            "Copy the UI, workflows, builder-site layouts, reports, dashboards, or design systems for competing or commercial use.",
            "Resell, sublicense, lease, scrape, frame, or commercially exploit the platform or protected parts of it.",
            "Remove or alter copyright, attribution, trademark, or ownership notices.",
        ],
    ),
    (
        "Branding and Trademarks",
        [
            '"Northstone", "Northstone CRM", "Deal Intelligence OS", related names, logos, marks, and brand assets are proprietary to Northstone CRM and may not be used without prior written permission except for legitimate nominative reference.'
        ],
    ),
    (
        "Databases and Structured Information",
        [
            "The selection, coordination, organization, formatting, and presentation of CRM data structures, database views, builder listing arrangements, and reporting layouts may also be protected to the maximum extent permitted by applicable intellectual property, database, and trade secret laws."
        ],
    ),
    (
        "Feedback",
        [
            "If any user submits suggestions, feature ideas, workflow recommendations, or product feedback relating to Northstone CRM, Northstone CRM may use such feedback without restriction or compensation unless otherwise agreed in writing."
        ],
    ),
    (
        "Third-Party Services",
        [
            "Northstone CRM may integrate with third-party services including cloud infrastructure, communication tools, workspace providers, AI providers, or payment systems. All third-party names and services remain the property of their respective owners. This notice covers Northstone CRM's own implementation, platform logic, and proprietary expression only."
        ],
    ),
    (
        "Reservation of Rights",
        [
            "All rights not expressly granted are reserved. No implied license, transfer, assignment, or waiver of intellectual property rights is created through access to or use of the platform."
        ],
    ),
    (
        "Contact for Copyright and IP Matters",
        [
            "Email: niharlakhani2@gmail.com",
            "Phone: +91 98342 41892",
        ],
    ),
]


def add_paragraph(doc: Document, text: str, *, bold: bool = False, size: int = 11) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Northstone CRM\nCopyright and Intellectual Property Notice")
    run.bold = True
    run.font.size = Pt(18)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        "Last updated: June 3, 2026\n"
        "Website: https://www.northstonecrm.com\n"
        "App: https://app.northstonecrm.com\n"
        "Contact: niharlakhani2@gmail.com | +91 98342 41892"
    )
    meta_run.font.size = Pt(10)

    intro = (
        "Northstone CRM, Deal Intelligence OS, and all related platform materials are proprietary. "
        "This notice explains ownership of the software, interfaces, builder website systems, data "
        "structures, workflows, generated assets, and supporting materials that make up the Northstone CRM platform."
    )
    add_paragraph(doc, intro, size=11)

    for index, (heading, items) in enumerate(SECTIONS, start=1):
        add_paragraph(doc, f"{index}. {heading}", bold=True, size=13)
        for item in items:
            if item.endswith(":"):
                add_paragraph(doc, item, size=11)
            elif heading in {"Materials Covered by Copyright", "Protected Platform Functionality", "Restrictions"} and item not in {
                items[0]
            }:
                paragraph = doc.add_paragraph(style=None)
                paragraph.style = doc.styles["List Bullet"]
                run = paragraph.add_run(item)
                run.font.size = Pt(11)
            else:
                add_paragraph(doc, item, size=11)

    add_paragraph(
        doc,
        "Short-form notice:",
        bold=True,
        size=12,
    )
    add_paragraph(
        doc,
        "© 2026 Northstone CRM. All rights reserved. Software, UI, builder website systems, templates, databases, documents, reports, and platform content are proprietary.",
        size=10,
    )

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
